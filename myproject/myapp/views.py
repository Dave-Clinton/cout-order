# users/views.py
from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout
from django.core.mail import send_mail
from django.conf import settings
from .models import Profile, Affidavit
from .forms import ReviewForm,AffidavitForm
import uuid
from django.utils.html import escape
from django.core.exceptions import ValidationError
from django.urls import reverse
from django.utils.http import urlsafe_base64_encode
from django.utils.encoding import force_bytes
from django.template.loader import render_to_string
from django.shortcuts import get_object_or_404
from django.utils.http import urlsafe_base64_decode
from django.contrib.auth.models import User
from django.http import HttpResponse
from django.utils.encoding import force_str
from django.core.validators import validate_email
from django.contrib.auth import authenticate, login as auth_login
from django.contrib.auth.tokens import default_token_generator
from django.shortcuts import render, redirect, get_object_or_404
from .models import Profile, Affidavit
from django.contrib.auth.decorators import user_passes_test
from django.http import HttpResponseForbidden





def register(request):
    if request.method == 'POST':
        first_name = request.POST['first_name']
        last_name = request.POST['last_name']
        email = request.POST['email']
        password = request.POST['password']
        password2 = request.POST['password2']

        if password == password2:
            if User.objects.filter(email=email).exists():
                messages.error(request, 'Email is already taken')
            elif len(password) < 8 or not any(char.isdigit() for char in password) or not any(char.isalpha() for char in password) or not any(char.islower() for char in password) or not any(char.isupper() for char in password):
                messages.error(request, 'Password must be at least 8 characters long, contain both numbers and letters, and have both lowercase and uppercase letters')
            else:
                is_admin = email in settings.ADMIN_EMAILS
                user = User.objects.create_user(username=email, email=email, password=password)  # Create a regular user
                user.save()
                verification_code = uuid.uuid4()
                profile = Profile(user=user, first_name=first_name, last_name=last_name, verification_code=verification_code)
                profile.is_admin = is_admin  # Assign admin status based on email
                profile.save()
                send_verification_email(user, verification_code)
                messages.success(request, 'Registration successful. Please check your email to verify your account.')
                return redirect('verify')
        else:
            messages.error(request, 'Passwords do not match')

    return render(request, 'register.html')




def send_verification_email(user, verification_code):
    uid = urlsafe_base64_encode(force_bytes(user.pk))
    token = str(verification_code)
    verification_link = reverse('verify_email', kwargs={'uidb64': uid, 'token': token})
    verification_url = f"{settings.SITE_URL}{verification_link}"

    subject = 'Verify your email address'
    message = f'Hi {user.username},\n\nPlease click the following link to verify your email address: {verification_url}\n\nThank you!'
    from_email = settings.DEFAULT_FROM_EMAIL
    recipient_list = [user.email]
    send_mail(subject, message, from_email, recipient_list)




def verify_email(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = get_object_or_404(User, pk=uid)
        profile = Profile.objects.get(user=user)
        verification_code = uuid.UUID(token)
        
        if profile.verification_code == verification_code and not profile.email_verified:
            profile.email_verified = True
            profile.save()
            messages.success(request, 'Email verified successfully!')
            return redirect('login')
        else:
            messages.error(request, 'Verification link is invalid or has already been used.')
            return render(request, 'verify.html')
    except (TypeError, ValueError, OverflowError, User.DoesNotExist, Profile.DoesNotExist):
        messages.error(request, 'Verification link is invalid.')
        return render(request, 'verify.html')


def verify (request):
    return render (request, 'verify.html')








def login_view(request):
    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '').strip()

        # Validate email
        try:
            validate_email(email)
        except ValidationError:
            messages.error(request, 'Invalid email format')
            return render(request, 'login.html')

        # Check if email and password are provided
        if not email or not password:
            messages.error(request, 'Email and password are required')
            return render(request, 'login.html')

        # Authenticate user
        user = authenticate(request, username=email, password=password)
        if user is not None:
            auth_login(request, user)
            return redirect('home')  # Adjust the redirect as needed
        else:
            messages.error(request, 'Invalid credentials')

    return render(request, 'login.html')





def password_reset_request(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        user = User.objects.filter(email=email).first()
        if user:
            uid = urlsafe_base64_encode(force_bytes(user.pk))
            token = default_token_generator.make_token(user)
            password_reset_link = reverse('password_reset_confirm', kwargs={'uidb64': uid, 'token': token})
            password_reset_url = f"{settings.SITE_URL}{password_reset_link}"

            subject = 'Reset your password'
            message = f'Hi {user.username},\n\nPlease click the following link to reset your password: {password_reset_url}\n\nThank you!'
            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = [user.email]
            send_mail(subject, message, from_email, recipient_list)

            messages.success(request, 'A link to reset your password has been sent to your email.')
            return redirect('login')
        else:
            messages.error(request, 'This email is not registered.')
    return render(request, 'password_reset_form.html')



def password_reset_confirm(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = get_object_or_404(User, pk=uid)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None

    if user is not None and default_token_generator.check_token(user, token):
        if request.method == 'POST':
            password = request.POST.get('password')
            password2 = request.POST.get('password2')
            if password == password2:
                if len(password) < 8 or not any(char.isdigit() for char in password) or not any(char.isalpha() for char in password) or not any(char.islower() for char in password) or not any(char.isupper() for char in password):
                    messages.error(request, 'Password must be at least 8 characters long, contain both numbers and letters, and have both lowercase and uppercase letters')
                else:
                    user.set_password(password)
                    user.save()
                    messages.success(request, 'Password has been reset successfully!')
                    return redirect('login')
            else:
                messages.error(request, 'Passwords do not match')
        return render(request, 'password_reset_confirm.html')
    else:
        messages.error(request, 'The reset link is invalid, possibly because it has already been used.')
        return redirect('password_reset_form')


from django.contrib.auth.models import User




@login_required
def home(request):
    user_profile = Profile.objects.get(user=request.user)
    completed_affidavits = Affidavit.objects.filter(profile=user_profile, status='complete')
    context = {
        'profile': user_profile,
        'completed_affidavits': completed_affidavits,
    }
    return render(request, 'home.html', context)



def logout_view(request):
    logout(request)
    return redirect('login')








@login_required
def file_new_case(request):
    if request.method == 'POST':
        user_profile = Profile.objects.get(user=request.user)
        user_profile.court_selection = request.POST.get('court-selection')
        user_profile.court_station = request.POST.get('court-station')
        user_profile.court_division = request.POST.get('court-division')
        user_profile.case_category = request.POST.get('case-category')
        user_profile.case_type = request.POST.get('case-type')
        user_profile.save()
        messages.success(request, 'Case filed successfully.')  
        return redirect('case_parties')
    return render(request, 'file_new_case.html')


@login_required
def case_parties(request):
    if request.method == 'POST':
        user_profile = Profile.objects.get(user=request.user)
        user_profile.party_type = request.POST.get('party-type')
        user_profile.party_level = request.POST.get('party-level')
        user_profile.case_party_type = request.POST.get('case-party-type')
        user_profile.organization_name = request.POST.get('organization-name')
        user_profile.kra_pin = request.POST.get('kra-pin')
        user_profile.postal_address = request.POST.get('postal-address')
        user_profile.physical_location = request.POST.get('physical-location')
        user_profile.mobile_number = request.POST.get('mobile-number')
        user_profile.organization_email = request.POST.get('organization-email')
        user_profile.save()
        messages.success(request, 'Case filed successfully.')  # Optional success message
        return redirect('details')
    return render(request, 'case_parties.html')

@login_required
def details(request):
    if request.method == 'POST':
        user_profile = Profile.objects.get(user=request.user)
        user_profile.tenant = request.POST.get('tenant')
        user_profile.postal_address = request.POST.get('postal_address')
        user_profile.telephone_number = request.POST.get('telephone_number')
        user_profile.landlord_name = request.POST.get('landlord_name')
        user_profile.agent = request.POST.get('agent')
        user_profile.caretaker = request.POST.get('caretaker')
        user_profile.auctioneer = request.POST.get('auctioneer')
        user_profile.duration_of_stay = request.POST.get('duration_of_stay')
        user_profile.monthly_rent = request.POST.get('monthly_rent')
        user_profile.year_of_entry = request.POST.get('year_of_entry')
        user_profile.deposit_paid = request.POST.get('deposit_paid')
        user_profile.cause_of_action = request.POST.get('cause_of_action')
        user_profile.problem = request.POST.get('problem')
        user_profile.ocs_police_station = request.POST.get('ocs_police_station')

        user_profile.organization_email = settings.CENTRAL_NOTIFICATION_EMAIL 
        user_profile.save()
        # Send an email to the user after saving the details
        send_mail(
            'Profile Details Saved',
            'Dear {},\n\nYour profile details have been successfully saved.\n\nThank you.'.format(user_profile.first_name),
            'fidelmasitsa03@gmail.com',  # From email
            [user_profile.user.email],  # To email
            fail_silently=False,
        )

        send_mail(
            'New Form Submission',
            f'Hello,\n\nA new form has been submitted by {request.user.username}.\n\nPlease check the details in the admin panel.',
            settings.EMAIL_HOST_USER,  # From email
            [settings.CENTRAL_NOTIFICATION_EMAIL],  # To admin email
            fail_silently=False,
        )

        return redirect('home')
        
    return render(request, 'details.html')





import requests
import base64
from datetime import datetime
import time
import json

def generate_access_token():
    consumer_key = "zi9kYKqrI2ZV2Xh8KHEodqK2yG8H46wNRz3nQcYpFLzGWZLp"
    consumer_secret = "439FtlzbXSAwo6vjG73eXvAVULR1konBjN4luaTjRAtSdPIDnosreddVe34KclAL"

    #live
    url = "https://api.safaricom.co.ke/oauth/v1/generate?grant_type=client_credentials"

    try:
        
        encoded_credentials = base64.b64encode(f"{consumer_key}:{consumer_secret}".encode()).decode()

        
        headers = {
            "Authorization": f"Basic {encoded_credentials}",
            "Content-Type": "application/json"
        }

        # Send the request and parse the response
        response = requests.get(url, headers=headers).json()

        # Check for errors and return the access token
        if "access_token" in response:
            return response["access_token"]
        else:
            raise Exception("Failed to get access token: " + response["error_description"])
    except Exception as e:
        raise Exception("Failed to get access token: " + str(e)) 



# View to render the HTML template
def mpesa_stk_push(request):
    return render(request, 'mpesa_stk_push.html')

 

def sendStkPush(phone_number, amount):
    token = generate_access_token()
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    shortCode = "5152644"

    passkey = "bfb279f9aa9bdbcf158e97dd71a467cd2e0c893059b10f78e6b72ada1ed2c919"
    stk_password = base64.b64encode((shortCode + passkey + timestamp).encode('utf-8')).decode('utf-8')

    url = "https://api.safaricom.co.ke/mpesa/stkpush/v1/processrequest"

    headers = {
        'Authorization': 'Bearer ' + token,
        'Content-Type': 'application/json'
    }

    requestBody = {
        "BusinessShortCode": shortCode,
        "Password": stk_password,
        "Timestamp": timestamp,
        "TransactionType": "CustomerBuyGoodsOnline",  # Or "CustomerBuyGoodsOnline"
        "Amount": amount,
        "PartyA": phone_number,
        "PartyB": 4565350,
        "PhoneNumber": phone_number,
        "CallBackURL": "https://yourwebsite.co.ke/callbackurl",
        "AccountReference": "account",
        "TransactionDesc": "test"
    }

    for attempt in range(3):  # Retry up to 3 times
        try:
            response = requests.post(url, json=requestBody, headers=headers).json()

            if response.get("ResponseCode") == "0":
                return {"status": "success", "data": response}
            elif response.get("errorCode") == "500.001.1001":
                time.sleep(5)  # Wait for 5 seconds before retrying
                continue
            else:
                return {"status": "error", "data": response}
        except Exception as e:
            return {"status": "error", "data": {'error': str(e)}}

    return {"status": "error", "data": {"errorMessage": "Max retries exceeded"}}


from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json

# Import the functions from your script
@csrf_exempt
def stk_push(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        phone_number = data.get('phoneNumber')
        amount = data.get('amount')

        try:
            # Call the sendStkPush function with the provided details
            response = sendStkPush(phone_number, amount)
            return JsonResponse(response)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid request'}, status=400)
@login_required
def dashboard(request):
    if request.user.email not in settings.ADMIN_EMAILS:
        return HttpResponseForbidden("You do not have permission to access this page.")
    profiles = Profile.objects.all()
    affidavits = Affidavit.objects.all()
    
    profiles_with_affidavits = [affidavit.profile.id for affidavit in affidavits]
    
    context = {
        'profiles': profiles,
        'affidavits': affidavits,
        'profiles_with_affidavits': profiles_with_affidavits
    }
    
    return render(request, 'dashboard.html', context)

@login_required
def add_affidavit(request, profile_id):
    profile = get_object_or_404(Profile, id=profile_id)
    
    # Redirect if affidavit already exists
    if Affidavit.objects.filter(profile=profile).exists():
        return redirect('dashboard')

    if request.method == 'POST':
        # Extract data from POST request
        content = request.POST.get('content')
        
        if content:
            # Create and save the new affidavit
            Affidavit.objects.create(profile=profile, content=content, status='active')
            return redirect('review_dashboard')
    
    # Render the template if GET request or invalid POST data
    return render(request, 'affidavit.html', {'profile': profile})

@login_required
def review_dashboard(request):
    profiles = Profile.objects.all()
    affidavits = Affidavit.objects.select_related('profile').all()

    if request.method == 'POST':
        affidavit_id = request.POST.get('affidavit_id')
        action = request.POST.get('action')
        affidavit = get_object_or_404(Affidavit, id=affidavit_id)

        if action == 'return':
            affidavit.status = 'in_revision'
            affidavit.save()
        elif action == 'complete':
            affidavit.status = 'complete'
            affidavit.save()

        # Refresh affidavits after status update
        affidavits = Affidavit.objects.select_related('profile').all()

    context = {
        'profiles': profiles,
        'affidavits': affidavits,
    }
    return render(request, 'review_dashboard.html', context)

@login_required
def edit_affidavit(request, affidavit_id):
    affidavit = get_object_or_404(Affidavit, id=affidavit_id)
    
    if request.method == 'POST':
        affidavit.content = request.POST.get('content')
        affidavit.status = 'in_review'
        affidavit.edited = True  
        affidavit.save()
        return redirect('dashboard')  
    
    return render(request, 'edit_affidavit.html', {'affidavit': affidavit})

@login_required
def delete_affidavit(request, affidavit_id):
    affidavit = get_object_or_404(Affidavit, id=affidavit_id)
    if request.method == 'POST':
        affidavit.delete()
        return redirect('dashboard')
    return render(request, 'confirm_delete.html', {'affidavit': affidavit})







#word
from django.http import HttpResponse
from docx import Document
from .models import Profile

def download_profiles_csv(request):
    # Create a new Document
    document = Document()

    # Add a title to the document
    document.add_heading('Profile List', level=1)

    # Add a table to the document
    table = document.add_table(rows=1, cols=28)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'First Name'
    hdr_cells[1].text = 'Last Name'
    hdr_cells[2].text = 'Court Selection'
    hdr_cells[3].text = 'Court Station'
    hdr_cells[4].text = 'Court Division'
    hdr_cells[5].text = 'Case Category'
    hdr_cells[6].text = 'Case Type'
    hdr_cells[7].text = 'Party Type'
    hdr_cells[8].text = 'Party Level'
    hdr_cells[9].text = 'Case Party Type'
    hdr_cells[10].text = 'Organization Name'
    hdr_cells[11].text = 'KRA Pin'
    hdr_cells[12].text = 'Postal Address'
    hdr_cells[13].text = 'Physical Location'
    hdr_cells[14].text = 'Mobile Number'
    hdr_cells[15].text = 'Organization Email'
    hdr_cells[16].text = 'Tenant'
    hdr_cells[17].text = 'Telephone Number'
    hdr_cells[18].text = 'Landlord Name'
    hdr_cells[19].text = 'Agent'
    hdr_cells[20].text = 'Caretaker'
    hdr_cells[21].text = 'Auctioneer'
    hdr_cells[22].text = 'Duration of Stay'
    hdr_cells[23].text = 'Monthly Rent'
    hdr_cells[24].text = 'Year of Entry'
    hdr_cells[25].text = 'Deposit Paid'
    hdr_cells[26].text = 'Cause of Action'
    hdr_cells[27].text = 'Problem'
    hdr_cells[28].text = 'OCS Police Station'

    # Add data rows to the table
    for profile in Profile.objects.all():
        row_cells = table.add_row().cells
        row_cells[0].text = profile.first_name
        row_cells[1].text = profile.last_name
        row_cells[2].text = profile.court_selection
        row_cells[3].text = profile.court_station
        row_cells[4].text = profile.court_division
        row_cells[5].text = profile.case_category
        row_cells[6].text = profile.case_type
        row_cells[7].text = profile.party_type
        row_cells[8].text = profile.party_level
        row_cells[9].text = profile.case_party_type
        row_cells[10].text = profile.organization_name
        row_cells[11].text = profile.kra_pin
        row_cells[12].text = profile.postal_address
        row_cells[13].text = profile.physical_location
        row_cells[14].text = profile.mobile_number
        row_cells[15].text = profile.organization_email
        row_cells[16].text = profile.tenant
        row_cells[17].text = profile.telephone_number
        row_cells[18].text = profile.landlord_name
        row_cells[19].text = profile.agent
        row_cells[20].text = profile.caretaker
        row_cells[21].text = profile.auctioneer
        row_cells[22].text = profile.duration_of_stay
        row_cells[23].text = profile.monthly_rent
        row_cells[24].text = profile.year_of_entry
        row_cells[25].text = profile.deposit_paid
        row_cells[26].text = profile.cause_of_action
        row_cells[27].text = profile.problem
        row_cells[28].text = profile.ocs_police_station

    # Create the HTTP response with Word content
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="profiles.docx"'

    # Save the document to the response
    document.save(response)

    return response



from django.http import HttpResponse
from docx import Document
from django.shortcuts import get_object_or_404
from .models import Profile

def download_profile_csv(request, profile_id):
    # Get the specific profile
    profile = get_object_or_404(Profile, id=profile_id)
    
    # Create a new Document
    document = Document()
    
    # Add a title to the document
    document.add_heading(f'Profile {profile_id}', level=1)
    
    # Add a table to the document
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    # Define fields and values
    fields = [
        ('Username', profile.user.username if profile.user else 'N/A'),
        ('First Name', profile.first_name or 'N/A'),
        ('Last Name', profile.last_name or 'N/A'),
        ('Court Selection', profile.court_selection or 'N/A'),
        ('Court Station', profile.court_station or 'N/A'),
        ('Court Division', profile.court_division or 'N/A'),
        ('Case Category', profile.case_category or 'N/A'),
        ('Case Type', profile.case_type or 'N/A'),
        ('Party Type', profile.party_type or 'N/A'),
        ('Party Level', profile.party_level or 'N/A'),
        ('Case Party Type', profile.case_party_type or 'N/A'),
        ('Organization Name', profile.organization_name or 'N/A'),
        ('KRA Pin', profile.kra_pin or 'N/A'),
        ('Postal Address', profile.postal_address or 'N/A'),
        ('Physical Location', profile.physical_location or 'N/A'),
        ('Mobile Number', profile.mobile_number or 'N/A'),
        ('Organization Email', profile.organization_email or 'N/A'),
        ('Tenant', profile.tenant or 'N/A'),
        ('Telephone Number', profile.telephone_number or 'N/A'),
        ('Landlord Name', profile.landlord_name or 'N/A'),
        ('Agent', profile.agent or 'N/A'),
        ('Caretaker', profile.caretaker or 'N/A'),
        ('Auctioneer', profile.auctioneer or 'N/A'),
        ('Duration of Stay', profile.duration_of_stay or 'N/A'),
        ('Monthly Rent', profile.monthly_rent or 'N/A'),
        ('Year of Entry', profile.year_of_entry or 'N/A'),
        ('Deposit Paid', profile.deposit_paid or 'N/A'),
        ('Cause of Action', profile.cause_of_action or 'N/A'),
        ('Problem', profile.problem or 'N/A'),
        ('OCS Police Station', profile.ocs_police_station or 'N/A')
    ]
    
    # Add data rows to the table
    for field, value in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = value
    
    # Create the HTTP response with Word content
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="profile_{profile_id}.docx"'
    
    # Save the document to the response
    document.save(response)
    
    return response

import csv

def download_completed_affidavits_csv(request):
    # Get the completed affidavits
    completed_affidavits = Affidavit.objects.filter(status='complete')
    
    # Create the HTTP response with CSV content
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="completed_affidavits.csv"'

    # Create a CSV writer object
    writer = csv.writer(response)
    
    # Write the header row
    writer.writerow(['Content', 'Status'])
    
    # Write the data rows
    for affidavit in completed_affidavits:
        writer.writerow([affidavit.content, affidavit.status.title()])

    return response




import csv
from django.http import HttpResponse
from django.template.loader import render_to_string
from docx import Document

def download_affidavits_csv(request, profile_id):
    # Get the affidavits for the specified profile
    affidavits = Affidavit.objects.filter(profile_id=profile_id)

    # Create a list of dictionaries for template rendering
    affidavit_data = [{
        'content': affidavit.content,
        'status': affidavit.get_status_display(),
        'edited': affidavit.edited
    } for affidavit in affidavits]

    # Render HTML table from template
    html_content = render_to_string('affidavit_table.html', {'affidavits': affidavit_data})

    # Create a Word document
    doc = Document()
    doc.add_paragraph()  # Add a blank paragraph for spacing
    doc.add_paragraph(html_content, style='Normal')

    # Save the document to a file-like object
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create the HTTP response
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="affidavits_profile_{profile_id}.docx"'

    return response











def download_affidavit_csv(request, affidavit_id):
    affidavit = get_object_or_404(Affidavit, id=affidavit_id)
    
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="affidavit_{affidavit_id}.csv"'

    writer = csv.writer(response)
    writer.writerow(['Content', 'Status', 'Edited'])
    writer.writerow([affidavit.content, affidavit.get_status_display(), affidavit.edited])

    return response










from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
import paypalrestsdk
import json

# Configure PayPal SDK
paypalrestsdk.configure({
    "mode": settings.PAYPAL_MODE,
    "client_id": settings.PAYPAL_CLIENT_ID,
    "client_secret": settings.PAYPAL_CLIENT_SECRET
})

def make_payment(request):
    # Render the template with the PayPal button
    return render(request, 'make_payment.html')

@csrf_exempt
def capture_payment(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        order_id = data.get('orderID')

        # Capture the payment
        payment = paypalrestsdk.Payment.find(order_id)
        if payment.execute({"payer_id": data.get('payerID')}):
            return JsonResponse({'status': 'success', 'message': 'Payment successfully captured'})
        else:
            return JsonResponse({'status': 'error', 'message': 'Payment capture failed'}, status=400)
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)




def index (request):

    return render (request, 'index.html')